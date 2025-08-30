# informes/views.py
from django.shortcuts import render, get_object_or_404, redirect
from django.contrib.auth.decorators import login_required
from django.http import HttpResponse
from .models import Informe
from .forms import InformeForm
from django.template.loader import render_to_string
from xhtml2pdf import pisa
from docx import Document
from io import BytesIO

# Plantillas pre-cargadas separadas por secciones
TIPO_1_PLANTILLA = {
    'motivo': "xxxx asiste a la consulta por iniciativa del equipo de ASM y sus padres...",
    'antecedentes': "Como antecedente significativo, la madre de xxxxx...",
    'sesiones': "- Sesión de motivo de consulta y anamnesis\n...",
    'resultados': "xxxxxxx se presenta de manera cooperativa durante la sesión..."
}

TIPO_2_PLANTILLA = {
    'motivo': "[Describa el motivo de consulta]",
    'antecedentes': "[Complete los antecedentes familiares y personales]",
    'sesiones': "[Liste sesiones e instrumentos aplicados]",
    'resultados': "[Describa resultados y sugerencias]"
}

# ----------------------------
# VISTAS PRINCIPALES
# ----------------------------

@login_required
def listar_informes(request):
    informes = Informe.objects.all().order_by('-fecha')
    return render(request, 'informes/listar_informes.html', {'informes': informes})

@login_required
def crear_informe(request):
    if request.method == 'POST':
        form = InformeForm(request.POST)
        if form.is_valid():
            informe = form.save(commit=False)
            # Combinar campos de secciones en contenido
            motivo = request.POST.get('motivo', '')
            antecedentes = request.POST.get('antecedentes', '')
            sesiones = request.POST.get('sesiones', '')
            resultados = request.POST.get('resultados', '')
            informe.contenido = f"II. Motivo de Consulta\n{motivo}\n\nIII. Antecedentes relevantes\n{antecedentes}\n\nIV. Sesiones e instrumentos aplicados\n{sesiones}\n\nV. Resultados y sugerencias\n{resultados}"
            # informe.autor = request.user  # si agregas campo autor
            informe.save()
            return redirect('detalle_informe', pk=informe.pk)
    else:
        form = InformeForm()
    return render(request, 'informes/crear_informe.html', {'form': form})

@login_required
def detalle_informe(request, pk):
    informe = get_object_or_404(Informe, pk=pk)
    return render(request, 'informes/detalle_informe.html', {'informe': informe})

@login_required
def editar_informe(request, pk):
    informe = get_object_or_404(Informe, pk=pk)
    if request.method == 'POST':
        form = InformeForm(request.POST, instance=informe)
        if form.is_valid():
            # Combinar campos de secciones en contenido
            motivo = request.POST.get('motivo', '')
            antecedentes = request.POST.get('antecedentes', '')
            sesiones = request.POST.get('sesiones', '')
            resultados = request.POST.get('resultados', '')
            informe.contenido = f"II. Motivo de Consulta\n{motivo}\n\nIII. Antecedentes relevantes\n{antecedentes}\n\nIV. Sesiones e instrumentos aplicados\n{sesiones}\n\nV. Resultados y sugerencias\n{resultados}"
            form.save()
            return redirect('detalle_informe', pk=informe.pk)
    else:
        form = InformeForm(instance=informe)
    return render(request, 'informes/editar_informe.html', {'form': form, 'informe': informe})

@login_required
def ver_informe(request, pk):
    informe = get_object_or_404(Informe, pk=pk)
    return render(request, 'informes/ver_informe.html', {'informe': informe})

# ----------------------------
# EXPORTACIÓN PDF (xhtml2pdf)
# ----------------------------
@login_required
def exportar_pdf(request, pk):
    informe = get_object_or_404(Informe, pk=pk)
    html = render_to_string('informes/pdf_template.html', {'informe': informe})
    result = BytesIO()
    pdf = pisa.pisaDocument(BytesIO(html.encode("UTF-8")), result)
    if not pdf.err:
        response = HttpResponse(result.getvalue(), content_type='application/pdf')
        response['Content-Disposition'] = f'attachment; filename=Informe_{informe.paciente.nombre}_{informe.fecha}.pdf'
        return response
    return HttpResponse("Error al generar PDF", status=400)

# ----------------------------
# EXPORTACIÓN WORD
# ----------------------------
@login_required
def exportar_word(request, pk):
    informe = get_object_or_404(Informe, pk=pk)
    document = Document()
    document.add_heading(f'Informe {informe.get_tipo_display()}', 0)
    document.add_paragraph(f'Paciente: {informe.paciente.nombre}')
    document.add_paragraph(f'Fecha: {informe.fecha}')
    document.add_paragraph('')
    document.add_paragraph(informe.contenido)
    
    f = BytesIO()
    document.save(f)
    f.seek(0)
    response = HttpResponse(f.read(), content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
    response['Content-Disposition'] = f'attachment; filename=Informe_{informe.paciente.nombre}_{informe.fecha}.docx'
    return response
